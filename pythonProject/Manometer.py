import tkinter as tk
from tkinter import ttk, messagebox
from openpyxl import load_workbook
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import random
from docx.enum.section import WD_ORIENT

# ======================== ГЛОБАЛЬНЫЕ НАСТРОЙКИ ========================
win = tk.Tk()
win.geometry("680x840")
win.title("Калибровка манометров")
win.configure(background='#21252b')

current_date = date.today()

# ======================== ФОНОВОЕ ИЗОБРАЖЕНИЕ ========================
bg_image = tk.PhotoImage(file="logo.png")
bg_label = ttk.Label(win, image=bg_image)
bg_label.place(x=0, y=0)


# ======================== ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ========================
def get_safe_float(value, default=0.0):
    """Безопасное преобразование строки в float"""
    try:
        return float(value)
    except (ValueError, TypeError):
        return default


def on_context_menu(event):
    """Обработка правого клика по вкладкам"""
    tab_id = notebook.identify(event.x, event.y)
    if tab_id:
        x = event.x
        if 10 <= x < 177:
            index = 0
            print(f'ПКМ: {notebook.tab(index)["text"]}')
        elif 177 <= x < 342:
            index = 1
            print(f'ПКМ: {notebook.tab(index)["text"]}')
        elif 342 <= x < 508:
            index = 2
            print(f'ПКМ: {notebook.tab(index)["text"]}')


# ======================== СОЗДАЁМ NOTEBOOK (КОНТЕЙНЕР ВКЛАДОК) ========================
notebook = ttk.Notebook(win, width=300, height=300)


# ======================== ВКЛАДКА 1: ДАННЫЕ МАНОМЕТРА ========================
tab_manometer_data = ttk.Frame(notebook)

# Фоновое изображение для вкладки
bg_tab1 = tk.PhotoImage(file="logo.png")
bg_label_tab1 = ttk.Label(tab_manometer_data, image=bg_tab1)
bg_label_tab1.place(x=0, y=0)

# --- Поля ввода ---
ttk.Label(tab_manometer_data, text="Название манометра").pack(padx=5, pady=3)
entry_name = ttk.Entry(tab_manometer_data)
entry_name.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Номер манометра").pack(padx=5, pady=3)
entry_serial_number = ttk.Entry(tab_manometer_data)
entry_serial_number.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Класс точности (допустимая погрешность, %)").pack(padx=5, pady=3)
entry_accuracy_class = ttk.Entry(tab_manometer_data)
entry_accuracy_class.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Диапазон измерений (верхний предел)").pack(padx=5, pady=3)
entry_range_max = ttk.Entry(tab_manometer_data)
entry_range_max.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Единицы измерения (кПа, МПа, кгс/см²)").pack(padx=5, pady=3)
entry_units = ttk.Entry(tab_manometer_data)
entry_units.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Применяемые средства калибровки").pack(padx=5, pady=3)
entry_calibration_tools = ttk.Entry(tab_manometer_data)
entry_calibration_tools.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Внешний осмотр").pack(padx=5, pady=3)
entry_external_inspection = ttk.Entry(tab_manometer_data)
entry_external_inspection.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="ФИО калибровщика").pack(padx=5, pady=3)
entry_calibrator_name = ttk.Entry(tab_manometer_data)
entry_calibrator_name.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Температура (°C)").pack(padx=5, pady=3)
entry_temperature = ttk.Entry(tab_manometer_data)
entry_temperature.insert(0, "20")
entry_temperature.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Атмосферное давление (мм рт.ст.)").pack(padx=5, pady=3)
entry_pressure = ttk.Entry(tab_manometer_data)
entry_pressure.insert(0, "760")
entry_pressure.pack(padx=5, pady=3)

ttk.Label(tab_manometer_data, text="Относительная влажность (%)").pack(padx=5, pady=3)
entry_humidity = ttk.Entry(tab_manometer_data)
entry_humidity.insert(0, "60")
entry_humidity.pack(padx=5, pady=3)


# ======================== ВКЛАДКА 2: РАСЧЁТ ПОГРЕШНОСТИ ========================
tab_calculation = ttk.Frame(notebook)

# Фоновое изображение
bg_tab2 = tk.PhotoImage(file="logo.png")
bg_label_tab2 = ttk.Label(tab_calculation, image=bg_tab2)
bg_label_tab2.place(x=0, y=0)

# Две колонки: прямой и обратный ход
frame_left = ttk.Frame(tab_calculation)
frame_left.pack(side="left", fill="both", expand=True, padx=10, pady=5)

frame_right = ttk.Frame(tab_calculation)
frame_right.pack(side="right", fill="both", expand=True, padx=10, pady=5)

# ========== ПРЯМОЙ ХОД (левая колонка) ==========
ttk.Label(frame_left, text="Прямой ход (нагрузка)", font=('Arial', 10, 'bold')).pack(pady=5)

# Точка 1
ttk.Label(frame_left, text="Оцифрованная точка 1 (эталон)").pack(padx=5, pady=3)
entry_point_1_direct = ttk.Entry(frame_left)
entry_point_1_direct.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания калибруемого прибора 1").pack(padx=5, pady=3)
entry_device_1_direct = ttk.Entry(frame_left)
entry_device_1_direct.pack(padx=5, pady=3)

# Точка 2
ttk.Label(frame_left, text="Оцифрованная точка 2 (эталон)").pack(padx=5, pady=3)
entry_point_2_direct = ttk.Entry(frame_left)
entry_point_2_direct.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания калибруемого прибора 2").pack(padx=5, pady=3)
entry_device_2_direct = ttk.Entry(frame_left)
entry_device_2_direct.pack(padx=5, pady=3)

# Точка 3
ttk.Label(frame_left, text="Оцифрованная точка 3 (эталон)").pack(padx=5, pady=3)
entry_point_3_direct = ttk.Entry(frame_left)
entry_point_3_direct.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания калибруемого прибора 3").pack(padx=5, pady=3)
entry_device_3_direct = ttk.Entry(frame_left)
entry_device_3_direct.pack(padx=5, pady=3)

# Точка 4
ttk.Label(frame_left, text="Оцифрованная точка 4 (эталон)").pack(padx=5, pady=3)
entry_point_4_direct = ttk.Entry(frame_left)
entry_point_4_direct.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания калибруемого прибора 4").pack(padx=5, pady=3)
entry_device_4_direct = ttk.Entry(frame_left)
entry_device_4_direct.pack(padx=5, pady=3)

# Точка 5
ttk.Label(frame_left, text="Оцифрованная точка 5 (эталон)").pack(padx=5, pady=3)
entry_point_5_direct = ttk.Entry(frame_left)
entry_point_5_direct.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания калибруемого прибора 5").pack(padx=5, pady=3)
entry_device_5_direct = ttk.Entry(frame_left)
entry_device_5_direct.pack(padx=5, pady=3)

# ========== ОБРАТНЫЙ ХОД (правая колонка) ==========
ttk.Label(frame_right, text="Обратный ход (разгрузка)", font=('Arial', 10, 'bold')).pack(pady=5)

# Точка 1
ttk.Label(frame_right, text="Оцифрованная точка 1 (эталон)").pack(padx=5, pady=3)
entry_point_1_reverse = ttk.Entry(frame_right)
entry_point_1_reverse.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания калибруемого прибора 1").pack(padx=5, pady=3)
entry_device_1_reverse = ttk.Entry(frame_right)
entry_device_1_reverse.pack(padx=5, pady=3)

# Точка 2
ttk.Label(frame_right, text="Оцифрованная точка 2 (эталон)").pack(padx=5, pady=3)
entry_point_2_reverse = ttk.Entry(frame_right)
entry_point_2_reverse.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания калибруемого прибора 2").pack(padx=5, pady=3)
entry_device_2_reverse = ttk.Entry(frame_right)
entry_device_2_reverse.pack(padx=5, pady=3)

# Точка 3
ttk.Label(frame_right, text="Оцифрованная точка 3 (эталон)").pack(padx=5, pady=3)
entry_point_3_reverse = ttk.Entry(frame_right)
entry_point_3_reverse.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания калибруемого прибора 3").pack(padx=5, pady=3)
entry_device_3_reverse = ttk.Entry(frame_right)
entry_device_3_reverse.pack(padx=5, pady=3)

# Точка 4
ttk.Label(frame_right, text="Оцифрованная точка 4 (эталон)").pack(padx=5, pady=3)
entry_point_4_reverse = ttk.Entry(frame_right)
entry_point_4_reverse.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания калибруемого прибора 4").pack(padx=5, pady=3)
entry_device_4_reverse = ttk.Entry(frame_right)
entry_device_4_reverse.pack(padx=5, pady=3)

# Точка 5
ttk.Label(frame_right, text="Оцифрованная точка 5 (эталон)").pack(padx=5, pady=3)
entry_point_5_reverse = ttk.Entry(frame_right)
entry_point_5_reverse.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания калибруемого прибора 5").pack(padx=5, pady=3)
entry_device_5_reverse = ttk.Entry(frame_right)
entry_device_5_reverse.pack(padx=5, pady=3)


# ======================== ВКЛАДКА 3: РЕЗУЛЬТАТЫ ========================
tab_results = ttk.Frame(notebook)

# Фоновое изображение
bg_tab3 = tk.PhotoImage(file="logo.png")
bg_label_tab3 = ttk.Label(tab_results, image=bg_tab3)
bg_label_tab3.place(x=0, y=0)

# Метки для отображения результатов прямого хода
ttk.Label(tab_results, text="Прямой ход (погрешность, %)", font=('Arial', 10, 'bold')).pack(pady=5)

label_error_1_direct = ttk.Label(tab_results, text='Точка 1: ')
label_error_1_direct.pack(padx=5, pady=2)

label_error_2_direct = ttk.Label(tab_results, text='Точка 2: ')
label_error_2_direct.pack(padx=5, pady=2)

label_error_3_direct = ttk.Label(tab_results, text='Точка 3: ')
label_error_3_direct.pack(padx=5, pady=2)

label_error_4_direct = ttk.Label(tab_results, text='Точка 4: ')
label_error_4_direct.pack(padx=5, pady=2)

label_error_5_direct = ttk.Label(tab_results, text='Точка 5: ')
label_error_5_direct.pack(padx=5, pady=2)

# Метки для отображения результатов обратного хода
ttk.Label(tab_results, text="Обратный ход (погрешность, %)", font=('Arial', 10, 'bold')).pack(pady=5)

label_error_1_reverse = ttk.Label(tab_results, text='Точка 1: ')
label_error_1_reverse.pack(padx=5, pady=2)

label_error_2_reverse = ttk.Label(tab_results, text='Точка 2: ')
label_error_2_reverse.pack(padx=5, pady=2)

label_error_3_reverse = ttk.Label(tab_results, text='Точка 3: ')
label_error_3_reverse.pack(padx=5, pady=2)

label_error_4_reverse = ttk.Label(tab_results, text='Точка 4: ')
label_error_4_reverse.pack(padx=5, pady=2)

label_error_5_reverse = ttk.Label(tab_results, text='Точка 5: ')
label_error_5_reverse.pack(padx=5, pady=2)


# ======================== ОСНОВНЫЕ ФУНКЦИИ РАСЧЁТА ========================
def calculate_errors():
    """Рассчитывает погрешность для всех точек (прямой и обратный ход)"""
    try:
        max_range = get_safe_float(entry_range_max.get())
        if max_range <= 0:
            messagebox.showerror("Ошибка", "Диапазон измерений должен быть больше 0")
            return None, None

        # Получаем все значения из полей ввода
        points_direct = [
            get_safe_float(entry_point_1_direct.get()),
            get_safe_float(entry_point_2_direct.get()),
            get_safe_float(entry_point_3_direct.get()),
            get_safe_float(entry_point_4_direct.get()),
            get_safe_float(entry_point_5_direct.get())
        ]

        devices_direct = [
            get_safe_float(entry_device_1_direct.get()),
            get_safe_float(entry_device_2_direct.get()),
            get_safe_float(entry_device_3_direct.get()),
            get_safe_float(entry_device_4_direct.get()),
            get_safe_float(entry_device_5_direct.get())
        ]

        points_reverse = [
            get_safe_float(entry_point_1_reverse.get()),
            get_safe_float(entry_point_2_reverse.get()),
            get_safe_float(entry_point_3_reverse.get()),
            get_safe_float(entry_point_4_reverse.get()),
            get_safe_float(entry_point_5_reverse.get())
        ]

        devices_reverse = [
            get_safe_float(entry_device_1_reverse.get()),
            get_safe_float(entry_device_2_reverse.get()),
            get_safe_float(entry_device_3_reverse.get()),
            get_safe_float(entry_device_4_reverse.get()),
            get_safe_float(entry_device_5_reverse.get())
        ]

        # Списки для хранения рассчитанных погрешностей
        errors_direct = []
        errors_reverse = []

        # Расчёт погрешности для прямого хода
        for i in range(5):
            deviation = devices_direct[i] - points_direct[i]
            error_percent = (deviation / max_range) * 100
            error_percent = round(error_percent, 3)
            errors_direct.append(error_percent)

        # Расчёт погрешности для обратного хода
        for i in range(5):
            deviation = devices_reverse[i] - points_reverse[i]
            error_percent = (deviation / max_range) * 100
            error_percent = round(error_percent, 3)
            errors_reverse.append(error_percent)

        # Обновляем метки на вкладке результатов
        error_labels_direct = [
            label_error_1_direct, label_error_2_direct,
            label_error_3_direct, label_error_4_direct, label_error_5_direct
        ]

        error_labels_reverse = [
            label_error_1_reverse, label_error_2_reverse,
            label_error_3_reverse, label_error_4_reverse, label_error_5_reverse
        ]

        for i, error in enumerate(errors_direct):
            error_labels_direct[i].config(text=f'Точка {i+1}: {error} %')

        for i, error in enumerate(errors_reverse):
            error_labels_reverse[i].config(text=f'Точка {i+1}: {error} %')

        # Активируем кнопку создания протокола
        btn_create_protocol.config(state='normal')

        return errors_direct, errors_reverse

    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при расчёте погрешности: {str(e)}")
        return None, None


def save_to_excel_journal():
    """Сохраняет данные калибровки в Excel-журнал"""
    try:
        name = entry_name.get()
        serial_number = entry_serial_number.get()
        accuracy_class = entry_accuracy_class.get()
        measurement_range = f"{entry_range_max.get()} {entry_units.get()}"
        calibrator_name = entry_calibrator_name.get()

        # Рассчитываем погрешности
        errors_direct, errors_reverse = calculate_errors()
        if errors_direct is None:
            return

        # Проверка на соответствие классу точности
        max_allowed = get_safe_float(accuracy_class)
        all_errors = errors_direct + errors_reverse
        is_valid = all(abs(e) <= max_allowed for e in all_errors)
        result = 'Годен' if is_valid else 'Не годен'

        # Путь к файлу Excel
        excel_file = 'Журнал.xlsx'

        # Проверяем существует ли файл, если нет — создаём
        if not os.path.exists(excel_file):
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = 'Лист1'
            ws.append(['Дата', 'Наименование СИ', 'Зав. номер', 'Класс точности',
                       'Диапазон', 'Результат', 'Калибровщик'])
            wb.save(excel_file)

        # Загружаем и добавляем данные
        wb = load_workbook(excel_file)
        ws = wb['Лист1']
        ws.append([
            current_date.strftime('%d.%m.%Y'),
            name,
            serial_number,
            accuracy_class,
            measurement_range,
            result,
            calibrator_name
        ])
        wb.save(excel_file)

        messagebox.showinfo("Успех", "Данные успешно сохранены в журнал")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при сохранении: {str(e)}")


def create_calibration_protocol():
    """Создаёт протокол калибровки в формате Word"""
    try:
        # Получаем данные из полей
        name = entry_name.get()
        serial_number = entry_serial_number.get()
        accuracy_class = get_safe_float(entry_accuracy_class.get())
        max_range = get_safe_float(entry_range_max.get())
        units = entry_units.get()
        calibration_tools = entry_calibration_tools.get()
        external_inspection = entry_external_inspection.get()
        calibrator_name = entry_calibrator_name.get()
        temperature = entry_temperature.get()
        pressure = entry_pressure.get()
        humidity = entry_humidity.get()

        # Рассчитываем погрешности
        errors_direct, errors_reverse = calculate_errors()
        if errors_direct is None:
            return

        # Получаем все значения из полей ввода для таблицы
        points_direct = [
            get_safe_float(entry_point_1_direct.get()),
            get_safe_float(entry_point_2_direct.get()),
            get_safe_float(entry_point_3_direct.get()),
            get_safe_float(entry_point_4_direct.get()),
            get_safe_float(entry_point_5_direct.get())
        ]

        devices_direct = [
            get_safe_float(entry_device_1_direct.get()),
            get_safe_float(entry_device_2_direct.get()),
            get_safe_float(entry_device_3_direct.get()),
            get_safe_float(entry_device_4_direct.get()),
            get_safe_float(entry_device_5_direct.get())
        ]

        points_reverse = [
            get_safe_float(entry_point_1_reverse.get()),
            get_safe_float(entry_point_2_reverse.get()),
            get_safe_float(entry_point_3_reverse.get()),
            get_safe_float(entry_point_4_reverse.get()),
            get_safe_float(entry_point_5_reverse.get())
        ]

        devices_reverse = [
            get_safe_float(entry_device_1_reverse.get()),
            get_safe_float(entry_device_2_reverse.get()),
            get_safe_float(entry_device_3_reverse.get()),
            get_safe_float(entry_device_4_reverse.get()),
            get_safe_float(entry_device_5_reverse.get())
        ]

        # Создаём документ
        doc = Document()

        # Устанавливаем альбомную ориентацию
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(8)

        # Заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(
            "Система калибровки средств измерений ПАО «Газпром»\n"
            "Общество с ограниченной ответственностью «Газпром энерго»\n\n"
            "(ООО «Газпром энерго»)\n"
            "Надымский филиал ООО «Газпром энерго»"
        )
        title_run.bold = True
        title_run.font.size = Pt(12)

        doc.add_paragraph("\n" + "=" * 170 + "\n")

        # Информация о протоколе
        reg_num = doc.add_paragraph()
        reg_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reg_num.add_run("Регистрационный номер в Реестре аккредитованных лиц № 090004")

        protocol_num = doc.add_paragraph()
        protocol_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        protocol_run = protocol_num.add_run(
            f"ПРОТОКОЛ КАЛИБРОВКИ СРЕДСТВ ИЗМЕРЕНИЙ №____ от {current_date.strftime('%d.%m.%Y')}г."
        )
        protocol_run.bold = True

        # Данные манометра
        doc.add_paragraph(
            f"Наименование, тип, модификация СИ: {name}, заводской (серийный) номер № {serial_number}"
        )
        doc.add_paragraph(
            f"Диапазон измерений: 0-{max_range} {units}, пределы основной погрешности: {accuracy_class}%, "
            f"вид калибровки: периодическая"
        )

        # Условия проведения
        doc.add_paragraph(
            f"Условия проведения калибровки: температура {temperature}°С; "
            f"атмосферное давление {pressure} мм рт.ст., "
            f"относительная влажность {humidity}%."
        )

        # Методика
        doc.add_paragraph("В соответствии с: МК 30-0007-2023")
        doc.add_paragraph(f"Применяемые средства калибровки: {calibration_tools}")
        doc.add_paragraph(f"Внешний осмотр: {external_inspection}")

        # Результаты измерений
        doc.add_paragraph("\nРезультаты измерений и определение основной погрешности СИ:")

        # Создаём таблицу
        table = doc.add_table(rows=2, cols=9)
        table.style = 'Table Grid'
        table.autofit = False

        # Настройка ширины столбцов
        col_widths = [0.8, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7, 1.0]
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)

        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Калибруемые точки диапазона"
        table.cell(0, 0).merge(table.cell(1, 0))

        hdr_cells[1].text = "Значение контролируемого параметра при прямом ходе"
        table.cell(0, 1).merge(table.cell(0, 2))

        hdr_cells[3].text = "Значение контролируемого параметра при обратном ходе"
        table.cell(0, 3).merge(table.cell(0, 4))

        hdr_cells[5].text = "Погрешность, %"
        table.cell(0, 5).merge(table.cell(0, 8))

        # Вторая строка заголовков
        sub_hdr = table.rows[1].cells
        sub_hdr[1].text = "Показания калибруемого СИ"
        sub_hdr[2].text = "Показания средств калибровки"
        sub_hdr[3].text = "Показания калибруемого СИ"
        sub_hdr[4].text = "Показания средств калибровки"
        sub_hdr[5].text = "Прямой ход"
        sub_hdr[6].text = "Обратный ход"
        sub_hdr[7].text = "Вариация"
        sub_hdr[8].text = "Допустимая погрешность"

        # Центрируем заголовки
        for row in table.rows[:2]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Заполняем данные
        all_errors = []

        for i in range(5):
            row_cells = table.add_row().cells

            # Калибруемая точка (берём из прямого хода)
            row_cells[0].text = str(points_direct[i])

            # Прямой ход
            row_cells[1].text = str(points_direct[i])  # Эталон
            row_cells[2].text = str(devices_direct[i])  # СИ

            # Обратный ход
            row_cells[3].text = str(points_reverse[i])  # Эталон
            row_cells[4].text = str(devices_reverse[i])  # СИ

            # Погрешность прямого хода
            row_cells[5].text = str(errors_direct[i])
            all_errors.append(abs(errors_direct[i]))

            # Погрешность обратного хода
            row_cells[6].text = str(errors_reverse[i])
            all_errors.append(abs(errors_reverse[i]))

            # Вариация (разница между показаниями прибора на прямом и обратном ходе)
            variation = abs(devices_direct[i] - devices_reverse[i])
            row_cells[7].text = str(round(variation, 3))

            # Допустимая погрешность
            row_cells[8].text = str(accuracy_class)

        # Объединяем ячейки "Допустимая погрешность" по вертикали
        for i in range(2, 7):
            table.cell(2, 8).merge(table.cell(i, 8))

        # Заключение
        max_allowed = accuracy_class
        is_valid = all(e <= max_allowed for e in all_errors)
        conclusion = "Заключение: Годен" if is_valid else "Заключение: Не годен"

        doc.add_paragraph("\n" + conclusion)
        doc.add_paragraph(f"Калибровщик: {calibrator_name}")

        # Сохраняем
        filename = f"Протокол калибровки {name} №{serial_number} от {current_date.strftime('%d.%m.%Y')}.docx"
        doc.save(filename)

        messagebox.showinfo("Успех", f"Протокол калибровки успешно создан: {filename}")
        os.startfile(filename)

    except Exception as e:
        messagebox.showerror("Ошибка", f"При создании протокола произошла ошибка: {str(e)}")


# ======================== КНОПКИ УПРАВЛЕНИЯ ========================
btn_calculate = ttk.Button(tab_results, text="📐 Рассчитать погрешность", command=calculate_errors)
btn_calculate.pack(padx=5, pady=10)

btn_save = ttk.Button(tab_results, text="💾 Сохранить в журнал", command=save_to_excel_journal)
btn_save.pack(padx=5, pady=5)

btn_create_protocol = ttk.Button(tab_results, text="📄 Создать протокол калибровки",
                                 command=create_calibration_protocol, state='disabled')
btn_create_protocol.pack(padx=5, pady=5)


# ======================== СБОРКА ИНТЕРФЕЙСА ========================
notebook.add(tab_manometer_data, text="Данные манометра")
notebook.add(tab_calculation, text="Расчёт погрешности")
notebook.add(tab_results, text="Результаты")

notebook.pack(fill="both", expand=1, padx=0, pady=0)
notebook.enable_traversal()
notebook.bind("<Button-3>", on_context_menu)

# Запуск
win.mainloop()
