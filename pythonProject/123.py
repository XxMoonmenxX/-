import tkinter as tk
from tkinter import ttk, messagebox
from openpyxl import load_workbook
import sqlite3 as sql
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx.enum.section import WD_ORIENT

con = sql.connect('test.txt')
cur = con.cursor()
soc = cur.fetchall()

win = tk.Tk()
win.geometry("680x840")  # Увеличил размер окна для новых полей
bg = tk.PhotoImage(file="logo.png")
img = ttk.Label(win, image=bg)
img.place(x=0, y=0)


def oncontextaction(event):
    name_of_x_y = nb.identify(event.x, event.y)
    if name_of_x_y:
        x = event.x
        if 10 <= x < 177:
            index = 0
            print(f'ПКМ:  {nb.tab(index)["text"]}; index = {index}')
        if 177 <= x < 342:
            index = 1
            print(f'ПКМ:  {nb.tab(index)["text"]}; index = {index}')
        if 342 <= x < 508:
            index = 2
            print(f'ПКМ:  {nb.tab(index)["text"]}; index = {index}')


def procent_govna():
    try:
        # Прямой ход
        hui = str(float(entry3.get()) - float((entry4.get())))
        hui = str(round(float(hui), 5))
        loh = str(float(hui) / float(entry2.get()) * 100)
        loh = (round(float(loh), 3))
        label["text"] = loh

        huii = str(float(entryl3.get()) - float((entryl4.get())))
        huii = str(round(float(huii), 5))
        lohh = str(float(huii) / float(entry2.get()) * 100)
        lohh = (round(float(lohh), 3))
        label2["text"] = lohh

        huiii = str(float(entryll3.get()) - float((entryll4.get())))
        huiii = str(round(float(huiii), 5))
        lohhh = str(float(huiii) / float(entry2.get()) * 100)
        lohhh = (round(float(lohhh), 3))
        label3["text"] = lohhh

        huiiii = str(float(entrylll3.get()) - float((entrylll4.get())))
        huiiii = str(round(float(huiiii), 5))
        lohhhh = str(float(huiiii) / float(entry2.get()) * 100)
        lohhhh = (round(float(lohhhh), 3))
        label4["text"] = lohhhh

        huiiiii = str(float(entryllll3.get()) - float((entryllll4.get())))
        huiiiii = str(round(float(huiiiii), 5))
        lohhhhh = str(float(huiiiii) / float(entry2.get()) * 100)
        lohhhhh = (round(float(lohhhhh), 3))
        label5["text"] = lohhhhh

        # Обратный ход
        hui_rev = str(float(entry3_rev.get()) - float((entry4_rev.get())))
        hui_rev = str(round(float(hui_rev), 5))
        loh_rev = str(float(hui_rev) / float(entry2.get()) * 100)
        loh_rev = (round(float(loh_rev), 3))
        label_rev["text"] = loh_rev

        huii_rev = str(float(entryl3_rev.get()) - float((entryl4_rev.get())))
        huii_rev = str(round(float(huii_rev), 5))
        lohh_rev = str(float(huii_rev) / float(entry2.get()) * 100)
        lohh_rev = (round(float(lohh_rev), 3))
        label2_rev["text"] = lohh_rev

        huiii_rev = str(float(entryll3_rev.get()) - float((entryll4_rev.get())))
        huiii_rev = str(round(float(huiii_rev), 5))
        lohhh_rev = str(float(huiii_rev) / float(entry2.get()) * 100)
        lohhh_rev = (round(float(lohhh_rev), 3))
        label3_rev["text"] = lohhh_rev

        huiiii_rev = str(float(entrylll3_rev.get()) - float((entrylll4_rev.get())))
        huiiii_rev = str(round(float(huiiii_rev), 5))
        lohhhh_rev = str(float(huiiii_rev) / float(entry2.get()) * 100)
        lohhhh_rev = (round(float(lohhhh_rev), 3))
        label4_rev["text"] = lohhhh_rev

        huiiiii_rev = str(float(entryllll3_rev.get()) - float((entryllll4_rev.get())))
        huiiiii_rev = str(round(float(huiiiii_rev), 5))
        lohhhhh_rev = str(float(huiiiii_rev) / float(entry2.get()) * 100)
        lohhhhh_rev = (round(float(lohhhhh_rev), 3))
        label5_rev["text"] = lohhhhh_rev

        # Активируем кнопку создания протокола после расчета
        create_protocol_btn['state'] = 'normal'

    except ValueError:
        messagebox.showerror("Ошибка", "Проверьте правильность введенных данных")


def create_calibration_protocol():
    try:
        doc = Document()

        # Устанавливаем альбомную ориентацию
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)

        # Настройка полей
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(8)

        # Заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Система калибровки средств измерений ПАО «Газпром»\n"
                                  "Общество с ограниченной ответственностью «Газпром энерго»\n\n"
                                  "(ООО «Газпром энерго»)\n"
                                  "Надымский филиал ООО «Газпром энерго»")
        title_run.bold = True
        title_run.font.size = Pt(12)

        doc.add_paragraph("\n" + "=" * 170 + "\n")

        # Информация о протоколе
        reg_num = doc.add_paragraph()
        reg_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reg_num.add_run("Регистрационный номер в Реестре аккредитованных лиц № 090004")

        protocol_num = doc.add_paragraph()
        protocol_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        protocol_run = protocol_num.add_run(
            f"ПРОТОКОЛ КАЛИБРОВКИ СРЕДСТВ ИЗМЕРЕНИЙ №____ от {current_date.strftime('%d.%m.%Y')}г.")
        protocol_run.bold = True

        # Данные манометра
        doc.add_paragraph(
            f"Наименование, тип, модификация СИ: {entry5.get()}, заводской (серийный) номер № {entry6.get()}")
        doc.add_paragraph(
            f"диапазон измерений: 0-{entry2.get()} {entry22.get()}, пределы основной погрешности: {entry.get()}, "
            f"вид калибровки: периодическая")

        # Условия проведения
        doc.add_paragraph(f"Условия проведения калибровки: температура {entry_temp.get()}°С; "
                          f"атмосферное давление {entry_pressure.get()} мм рт.ст., "
                          f"относительная влажность {entry_humidity.get()}%.")

        # Методика калибровки
        doc.add_paragraph("В соответствии с: МК 30-0007-2023")
        doc.add_paragraph(f"Применяемые средства калибровки: {entry_calibration_tools.get()}")
        doc.add_paragraph(f"Внешний осмотр: {entry_inspection.get()}")

        # Результаты измерений
        doc.add_paragraph("\nРезультаты измерений и определение основной погрешности СИ:")

        # Создаем таблицу с 9 колонками
        table = doc.add_table(rows=2, cols=9)
        table.style = 'Table Grid'
        table.autofit = False

        # Настройка ширины столбцов
        col_widths = [0.8, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7, 1.0]
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width)

        # === Первая строка заголовков ===
        hdr_cells = table.rows[0].cells

        # 1. "Калибруемые точки" (вертикальное объединение 2 строк)
        hdr_cells[0].text = "Калибруемые точки диапазона"
        table.cell(0, 0).merge(table.cell(1, 0))

        # 2. "Значение контролируемого параметра при прямом ходе" (ячейки 1-2)
        hdr_cells[1].text = "Значение контролируемого параметра при прямом ходе"
        table.cell(0, 1).merge(table.cell(0, 2))

        # 3. "Значение контролируемого параметра при обратном ходе" (ячейки 3-4)
        hdr_cells[3].text = "Значение контролируемого параметра при обратном ходе"
        table.cell(0, 3).merge(table.cell(0, 4))

        # 4. "Погрешность, %" (ячейки 5-8)
        hdr_cells[5].text = "Погрешность, %"
        table.cell(0, 5).merge(table.cell(0, 8))

        # === Вторая строка заголовков ===
        sub_hdr = table.rows[1].cells

        # Подписи для колонок
        sub_hdr[1].text = "Показания калибруемого СИ"
        sub_hdr[2].text = "Показания средств калибровки"

        sub_hdr[3].text = "Показания калибруемого СИ"
        sub_hdr[4].text = "Показания средств калибровки"

        # Подписи для столбцов погрешности
        sub_hdr[5].text = "Прямой ход"
        sub_hdr[6].text = "Обратный ход"
        sub_hdr[7].text = "Вариация"
        sub_hdr[8].text = "Допустимая погрешность"

        # Центрируем весь текст в заголовках
        for row in table.rows[:9]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # === Заполнение данных ===
        def safe_get(value, default="0.0"):
            try:
                return str(float(value))
            except (ValueError, TypeError):
                return default

        # Получаем данные из полей ввода
        direct_readings = [
            safe_get(entry4.get()),  # Точка 1 (прямой ход)
            safe_get(entryl4.get()),  # Точка 2
            safe_get(entryll4.get()),  # Точка 3
            safe_get(entrylll4.get()),  # Точка 4
            safe_get(entryllll4.get())  # Точка 5
        ]

        points_direct = [
            (safe_get(entry3.get()), direct_readings[0], safe_get(label["text"])),
            (safe_get(entryl3.get()), direct_readings[1], safe_get(label2["text"])),
            (safe_get(entryll3.get()), direct_readings[2], safe_get(label3["text"])),
            (safe_get(entrylll3.get()), direct_readings[3], safe_get(label4["text"])),
            (safe_get(entryllll3.get()), direct_readings[4], safe_get(label5["text"]))
        ]

        points_reverse = [
            (safe_get(entry3_rev.get()), safe_get(entry4_rev.get()), safe_get(label_rev["text"])),
            (safe_get(entryl3_rev.get()), safe_get(entryl4_rev.get()), safe_get(label2_rev["text"])),
            (safe_get(entryll3_rev.get()), safe_get(entryll4_rev.get()), safe_get(label3_rev["text"])),
            (safe_get(entrylll3_rev.get()), safe_get(entrylll4_rev.get()), safe_get(label4_rev["text"])),
            (safe_get(entryllll3_rev.get()), safe_get(entryllll4_rev.get()), safe_get(label5_rev["text"]))
        ]

        # Проверяем, соответствует ли манометр требованиям
        max_allowed_error = float(entry.get())
        passed = True  # Флаг соответствия требованиям

        # Список всех погрешностей для проверки
        all_errors = []

        # Заполняем данные по точкам
        for i in range(5):
            row_cells = table.add_row().cells

            # Вместо "Точка N" используем показания калибруемого СИ при прямом ходе
            row_cells[0].text = points_reverse[i][0]

            # Прямой ход
            row_cells[1].text = points_direct[i][0]  # Эталон
            row_cells[2].text = points_direct[i][1]  # СИ

            # Обратный ход
            row_cells[3].text = points_reverse[i][0]  # Эталон
            row_cells[4].text = points_reverse[i][1]  # СИ

            # Погрешность прямого хода
            direct_error = float(points_direct[i][2])
            row_cells[5].text = points_direct[i][2]  # Погрешность прямого хода
            all_errors.append(direct_error)

            # Погрешность обратного хода
            rev_error = ""
            if i == 0:
                rev_error = safe_get(label_rev["text"])
            elif i == 1:
                rev_error = safe_get(label2_rev["text"])
            elif i == 2:
                rev_error = safe_get(label3_rev["text"])
            elif i == 3:
                rev_error = safe_get(label4_rev["text"])
            else:
                rev_error = safe_get(label5_rev["text"])

            rev_error_value = float(rev_error)
            row_cells[6].text = rev_error
            all_errors.append(rev_error_value)

            # Вариация (разница между прямым и обратным ходом)
            variation = abs(float(points_direct[i][1]) - float(points_reverse[i][1]))
            row_cells[7].text = str(round(variation, 3))

            # Допустимая погрешность (одинаковая для всех строк)
            row_cells[8].text = entry.get()

        # Проверяем все погрешности
        for error in all_errors:
            if abs(error) > max_allowed_error:
                passed = False
                break

        # Объединяем ячейки "Допустимая погрешность" по вертикали
        for i in range(2, 7):  # Объединяем со 2 по 6 строку (индексы 1-5)
            table.cell(2, 8).merge(table.cell(i, 8))

        # Заключение на основе проверки
        if passed:
            conclusion_text = "Заключение: Годен"
        else:
            conclusion_text = "Заключение: Не годен"

        doc.add_paragraph("\n" + conclusion_text)
        doc.add_paragraph(f"Калибровщик: {entryy22.get()}")

        # Сохранение
        filename = f"Протокол калибровки {entry5.get()} №{entry6.get()} от {current_date.strftime('%d.%m.%Y')}.docx"
        doc.save(filename)

        messagebox.showinfo("Успех", f"Протокол калибровки успешно создан: {filename}")
        os.startfile(filename)

    except Exception as e:
        messagebox.showerror("Ошибка", f"При создании протокола произошла ошибка: {str(e)}")


def centrtxt():
    try:
        with con:
            print('Данные внесены ')
            name = entry5.get()
            number = str(entry6.get())
            kt = str(entry.get())
            imya = str(entryy22.get())
            diap = str(entry2.get() + " " + entry22.get())

            # Прямой ход
            hui = str(float(entry3.get()) - float((entry4.get())))
            hui = str(round(float(hui), 5))
            loh = str(float(hui) / float(entry2.get()) * 100)
            loh = (round(float(loh), 3))

            huii = str(float(entryl3.get()) - float((entryl4.get())))
            huii = str(round(float(huii), 5))
            lohh = str(float(huii) / float(entry2.get()) * 100)
            lohh = (round(float(lohh), 3))

            huiii = str(float(entryll3.get()) - float((entryll4.get())))
            huiii = str(round(float(huiii), 5))
            lohhh = str(float(huiii) / float(entry2.get()) * 100)
            lohhh = (round(float(lohhh), 3))

            huiiii = str(float(entrylll3.get()) - float((entrylll4.get())))
            huiiii = str(round(float(huiiii), 5))
            lohhhh = str(float(huiiii) / float(entry2.get()) * 100)
            lohhhh = (round(float(lohhhh), 3))

            huiiiii = str(float(entryllll3.get()) - float((entryllll4.get())))
            huiiiii = str(round(float(huiiiii), 5))
            lohhhhh = str(float(huiiiii) / float(entry2.get()) * 100)
            lohhhhh = (round(float(lohhhhh), 3))

            # Обратный ход
            hui_rev = str(float(entry3_rev.get()) - float((entry4_rev.get())))
            hui_rev = str(round(float(hui_rev), 5))
            loh_rev = str(float(hui_rev) / float(entry2.get()) * 100)
            loh_rev = (round(float(loh_rev), 3))

            huii_rev = str(float(entryl3_rev.get()) - float((entryl4_rev.get())))
            huii_rev = str(round(float(huii_rev), 5))
            lohh_rev = str(float(huii_rev) / float(entry2.get()) * 100)
            lohh_rev = (round(float(lohh_rev), 3))

            huiii_rev = str(float(entryll3_rev.get()) - float((entryll4_rev.get())))
            huiii_rev = str(round(float(huiii_rev), 5))
            lohhh_rev = str(float(huiii_rev) / float(entry2.get()) * 100)
            lohhh_rev = (round(float(lohhh_rev), 3))

            huiiii_rev = str(float(entrylll3_rev.get()) - float((entrylll4_rev.get())))
            huiiii_rev = str(round(float(huiiii_rev), 5))
            lohhhh_rev = str(float(huiiii_rev) / float(entry2.get()) * 100)
            lohhhh_rev = (round(float(lohhhh_rev), 3))

            huiiiii_rev = str(float(entryllll3_rev.get()) - float((entryllll4_rev.get())))
            huiiiii_rev = str(round(float(huiiiii_rev), 5))
            lohhhhh_rev = str(float(huiiiii_rev) / float(entry2.get()) * 100)
            lohhhhh_rev = (round(float(lohhhhh_rev), 3))

            # Проверка на соответствие классу точности
            if (float(entry.get()) >= abs(float(loh)) and float(entry.get()) >= abs(float(lohh)) and
                    float(entry.get()) >= abs(float(lohhh)) and float(entry.get()) >= abs(float(lohhhh)) and
                    float(entry.get()) >= abs(float(lohhhhh)) and
                    float(entry.get()) >= abs(float(loh_rev)) and float(entry.get()) >= abs(float(lohh_rev)) and
                    float(entry.get()) >= abs(float(lohhh_rev)) and float(entry.get()) >= abs(float(lohhhh_rev)) and
                    float(entry.get()) >= abs(float(lohhhhh_rev))):
                k = 'Годен'
            else:
                k = 'Не годен'

            xl = 'Журнал.xlsx'
            omg = load_workbook(xl)
            ogm = omg['Лист1']
            ogm.append([current_date.strftime('%d.%m.%Y'), name, number, kt, diap, k, imya])
            omg.save(xl)
            omg.close()

            messagebox.showinfo("Успех", "Данные успешно сохранены в журнал")

    except Exception as e:
        messagebox.showerror("Ошибка", f"При сохранении данных произошла ошибка: {str(e)}")


color = '#21252b'
win.configure(background=color)
current_date = date.today()

nb = ttk.Notebook(win, width=300, height=300)

fr1 = ttk.Frame(nb)
fr2 = ttk.Frame(nb)
fr3 = ttk.Frame(nb)

# Вкладка 1 - Данные манометра
bg = tk.PhotoImage(file="logo.png")
img = ttk.Label(fr1, image=bg)
img.place(x=0, y=0)

lb1 = ttk.Label(fr1, text="Название манометра")
lb1.pack(padx=5, pady=3)
entry5 = ttk.Entry(fr1)
entry5.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Номер манометра")
lb1.pack(padx=5, pady=3)
entry6 = ttk.Entry(fr1)
entry6.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Класс точности")
lb1.pack(padx=5, pady=3)
entry = ttk.Entry(fr1)
entry.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Диапазон")
lb1.pack(padx=5, pady=3)
entry2 = ttk.Entry(fr1)
entry2.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Единицы измерения (кПа, МПа, кгс/см²)")
lb1.pack(padx=5, pady=3)
entry22 = ttk.Entry(fr1)
entry22.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Применяемые средства калибровки")
lb1.pack(padx=5, pady=3)
entry_calibration_tools = ttk.Entry(fr1)
entry_calibration_tools.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Внешний осмотр")
lb1.pack(padx=5, pady=3)
entry_inspection = ttk.Entry(fr1)
entry_inspection.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="ФИО калибровщика")
lb1.pack(padx=5, pady=3)
entryy22 = ttk.Entry(fr1)
entryy22.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Температура (°C)")
lb1.pack(padx=5, pady=3)
entry_temp = ttk.Entry(fr1)
entry_temp.insert(0, "20")  # Значение по умолчанию
entry_temp.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Атмосферное давление (мм рт.ст.)")
lb1.pack(padx=5, pady=3)
entry_pressure = ttk.Entry(fr1)
entry_pressure.insert(0, "760")  # Значение по умолчанию
entry_pressure.pack(padx=5, pady=3)

lb1 = ttk.Label(fr1, text="Относительная влажность (%)")
lb1.pack(padx=5, pady=3)
entry_humidity = ttk.Entry(fr1)
entry_humidity.insert(0, "60")  # Значение по умолчанию
entry_humidity.pack(padx=5, pady=3)

# Вкладка 2 - Рассчет погрешности (переработанная версия с двумя колонками)
bg1 = tk.PhotoImage(file="logo.png")
img1 = ttk.Label(fr2, image=bg1)
img1.place(x=0, y=0)

# Создаем фреймы для двух колонок
frame_left = ttk.Frame(fr2)
frame_left.pack(side="left", fill="both", expand=True, padx=10, pady=5)

frame_right = ttk.Frame(fr2)
frame_right.pack(side="right", fill="both", expand=True, padx=10, pady=5)

# Прямой ход (левая колонка)
ttk.Label(frame_left, text="Прямой ход", font=('Arial', 10, 'bold')).pack(pady=5)

ttk.Label(frame_left, text="Оцифрованная точка 1").pack(padx=5, pady=3)
entry3 = ttk.Entry(frame_left)
entry3.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания эталона 1").pack(padx=5, pady=3)
entry4 = ttk.Entry(frame_left)
entry4.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Оцифрованная точка 2").pack(padx=5, pady=3)
entryl3 = ttk.Entry(frame_left)
entryl3.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания эталона 2").pack(padx=5, pady=3)
entryl4 = ttk.Entry(frame_left)
entryl4.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Оцифрованная точка 3").pack(padx=5, pady=3)
entryll3 = ttk.Entry(frame_left)
entryll3.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания эталона 3").pack(padx=5, pady=3)
entryll4 = ttk.Entry(frame_left)
entryll4.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Оцифрованная точка 4").pack(padx=5, pady=3)
entrylll3 = ttk.Entry(frame_left)
entrylll3.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания эталона 4").pack(padx=5, pady=3)
entrylll4 = ttk.Entry(frame_left)
entrylll4.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Оцифрованная точка 5").pack(padx=5, pady=3)
entryllll3 = ttk.Entry(frame_left)
entryllll3.pack(padx=5, pady=3)

ttk.Label(frame_left, text="Показания эталона 5").pack(padx=5, pady=3)
entryllll4 = ttk.Entry(frame_left)
entryllll4.pack(padx=5, pady=3)

# Обратный ход (правая колонка)
ttk.Label(frame_right, text="Обратный ход", font=('Arial', 10, 'bold')).pack(pady=5)

ttk.Label(frame_right, text="Оцифрованная точка 1").pack(padx=5, pady=3)
entry3_rev = ttk.Entry(frame_right)
entry3_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания эталона 1").pack(padx=5, pady=3)
entry4_rev = ttk.Entry(frame_right)
entry4_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Оцифрованная точка 2").pack(padx=5, pady=3)
entryl3_rev = ttk.Entry(frame_right)
entryl3_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания эталона 2").pack(padx=5, pady=3)
entryl4_rev = ttk.Entry(frame_right)
entryl4_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Оцифрованная точка 3").pack(padx=5, pady=3)
entryll3_rev = ttk.Entry(frame_right)
entryll3_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания эталона 3").pack(padx=5, pady=3)
entryll4_rev = ttk.Entry(frame_right)
entryll4_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Оцифрованная точка 4").pack(padx=5, pady=3)
entrylll3_rev = ttk.Entry(frame_right)
entrylll3_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания эталона 4").pack(padx=5, pady=3)
entrylll4_rev = ttk.Entry(frame_right)
entrylll4_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Оцифрованная точка 5").pack(padx=5, pady=3)
entryllll3_rev = ttk.Entry(frame_right)
entryllll3_rev.pack(padx=5, pady=3)

ttk.Label(frame_right, text="Показания эталона 5").pack(padx=5, pady=3)
entryllll4_rev = ttk.Entry(frame_right)
entryllll4_rev.pack(padx=5, pady=3)

fill_readings_btn = ttk.Button(frame_left, text="Заполнить показания", command=lambda: fill_readings())
fill_readings_btn.pack(padx=5, pady=10)


# Новая функция для автоматического заполнения показаний
def fill_readings():
    try:
        # Получаем погрешность и диапазон из полей ввода
        error = float(entry.get())
        range_val = float(entry2.get())

        # Вычисляем шаг для оцифрованных точек (5 точек равномерно по диапазону)
        step = range_val / 5

        # Заполняем оцифрованные точки для прямого хода
        entry3.delete(0, tk.END)
        entry3.insert(0, str(round(step * 1, 3)))
        entryl3.delete(0, tk.END)
        entryl3.insert(0, str(round(step * 2, 3)))
        entryll3.delete(0, tk.END)
        entryll3.insert(0, str(round(step * 3, 3)))
        entrylll3.delete(0, tk.END)
        entrylll3.insert(0, str(round(step * 4, 3)))
        entryllll3.delete(0, tk.END)
        entryllll3.insert(0, str(round(step * 5, 3)))

        # Заполняем оцифрованные точки для обратного хода
        entry3_rev.delete(0, tk.END)
        entry3_rev.insert(0, str(round(step * 1, 3)))
        entryl3_rev.delete(0, tk.END)
        entryl3_rev.insert(0, str(round(step * 2, 3)))
        entryll3_rev.delete(0, tk.END)
        entryll3_rev.insert(0, str(round(step * 3, 3)))
        entrylll3_rev.delete(0, tk.END)
        entrylll3_rev.insert(0, str(round(step * 4, 3)))
        entryllll3_rev.delete(0, tk.END)
        entryllll3_rev.insert(0, str(round(step * 5, 3)))

        # Функция для генерации случайного значения в пределах погрешности
        def get_random_value(base_value, error_percent):
            error_amount = base_value * error_percent / 100
            min_val = base_value - error_amount
            max_val = base_value + error_amount
            return round(random.uniform(min_val, max_val), 3)

        import random  # Добавляем в начало файла

        # Заполняем показания эталона для прямого хода (в пределах погрешности)
        base_val = float(entry3.get())
        entry4.delete(0, tk.END)
        entry4.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entryl3.get())
        entryl4.delete(0, tk.END)
        entryl4.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entryll3.get())
        entryll4.delete(0, tk.END)
        entryll4.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entrylll3.get())
        entrylll4.delete(0, tk.END)
        entrylll4.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entryllll3.get())
        entryllll4.delete(0, tk.END)
        entryllll4.insert(0, str(get_random_value(base_val, error)))

        # Заполняем показания эталона для обратного хода (в пределах погрешности)
        base_val = float(entry3_rev.get())
        entry4_rev.delete(0, tk.END)
        entry4_rev.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entryl3_rev.get())
        entryl4_rev.delete(0, tk.END)
        entryl4_rev.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entryll3_rev.get())
        entryll4_rev.delete(0, tk.END)
        entryll4_rev.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entrylll3_rev.get())
        entrylll4_rev.delete(0, tk.END)
        entrylll4_rev.insert(0, str(get_random_value(base_val, error)))

        base_val = float(entryllll3_rev.get())
        entryllll4_rev.delete(0, tk.END)
        entryllll4_rev.insert(0, str(get_random_value(base_val, error)))

        messagebox.showinfo("Успех", "Показания успешно заполнены автоматически")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при заполнении показаний: {str(e)}")

# Вкладка 3 - Результаты
bg2 = tk.PhotoImage(file="logo.png")
img2 = ttk.Label(fr3, image=bg2)
img2.place(x=0, y=0)

# Прямой ход
ttk.Label(fr3, text="Прямой ход", font=('Arial', 10, 'bold')).pack(pady=5)

label = ttk.Label(fr3, text='Точка 1: ')
label.pack(padx=5, pady=2)
label2 = ttk.Label(fr3, text='Точка 2: ')
label2.pack(padx=5, pady=2)
label3 = ttk.Label(fr3, text='Точка 3: ')
label3.pack(padx=5, pady=2)
label4 = ttk.Label(fr3, text='Точка 4: ')
label4.pack(padx=5, pady=2)
label5 = ttk.Label(fr3, text='Точка 5: ')
label5.pack(padx=5, pady=2)

# Обратный ход
ttk.Label(fr3, text="Обратный ход", font=('Arial', 10, 'bold')).pack(pady=5)

label_rev = ttk.Label(fr3, text='Точка 1: ')
label_rev.pack(padx=5, pady=2)
label2_rev = ttk.Label(fr3, text='Точка 2: ')
label2_rev.pack(padx=5, pady=2)
label3_rev = ttk.Label(fr3, text='Точка 3: ')
label3_rev.pack(padx=5, pady=2)
label4_rev = ttk.Label(fr3, text='Точка 4: ')
label4_rev.pack(padx=5, pady=2)
label5_rev = ttk.Label(fr3, text='Точка 5: ')
label5_rev.pack(padx=5, pady=2)

# Кнопки
btn = ttk.Button(fr3, text="Рассчет % погрешности шага", command=procent_govna)
btn.pack(padx=5, pady=10)

centrtext = ttk.Button(fr3, text="Внести данные", command=centrtxt)
centrtext.pack(padx=5, pady=5)

create_protocol_btn = ttk.Button(fr3, text="Создать протокол калибровки",
                                 command=create_calibration_protocol, state='disabled')
create_protocol_btn.pack(padx=5, pady=5)

# Добавление вкладок
nb.add(fr1, text="Данные манометра")
nb.add(fr2, text="Рассчет погрешности")
nb.add(fr3, text="Результаты")

nb.pack(fill="both", expand=1, padx=0, pady=0)
nb.enable_traversal()
nb.bind("<Button-3>", oncontextaction)

win.mainloop()
